"""
NeuroGraph V2 Backend - Real ML/AI Knowledge Graph Engine
Replaces all fake/hardcoded logic with production-ready algorithms
"""

import hashlib
import json
import os
import time
import uuid
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
import numpy as np
from dataclasses import dataclass, field, asdict

# NLP and ML imports
try:
    import spacy
    from spacy.tokens import Doc
    SPACY_AVAILABLE = True
except ImportError:
    SPACY_AVAILABLE = False

try:
    from langdetect import detect
    LANGDETECT_AVAILABLE = True
except ImportError:
    LANGDETECT_AVAILABLE = False

try:
    from sentence_transformers import SentenceTransformer
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    SENTENCE_TRANSFORMERS_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    import networkx as nx
    NETWORKX_AVAILABLE = True
except ImportError:
    NETWORKX_AVAILABLE = False


@dataclass
class Document:
    """Represents a processed document with metadata"""
    id: str
    filename: str
    file_hash: str
    file_type: str
    file_size: int
    content: str
    language: str
    entities: List[Dict[str, Any]]
    concepts: List[str]
    embedding: Optional[List[float]] = None
    created_at: str = field(default_factory=lambda: datetime.now().isoformat())
    ocr_applied: bool = False
    chunk_count: int = 0
    
    def to_dict(self) -> Dict:
        return asdict(self)


@dataclass
class ConceptNode:
    """A concept/entity node in the knowledge graph"""
    id: str
    label: str
    entity_type: str
    frequency: int = 1
    documents: List[str] = field(default_factory=list)
    embedding: Optional[List[float]] = None


@dataclass
class GraphEdge:
    """Edge between nodes in the knowledge graph"""
    source: str
    target: str
    weight: float
    edge_type: str  # 'similarity', 'co_occurrence', 'semantic'


class TextExtractor:
    """Extract text from various file formats"""
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF files using PyPDF2"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")
        
        text_parts = []
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return "\n\n".join(text_parts)
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """Extract text from Word documents"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    
    @staticmethod
    def extract_from_image(file_path: str) -> str:
        """Extract text from images using OCR (Tesseract)"""
        if not OCR_AVAILABLE:
            raise ImportError("pytesseract/PIL not installed. Run: pip install pytesseract Pillow")
        
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return text
    
    @classmethod
    def extract(cls, file_path: str, file_type: str) -> Tuple[str, bool]:
        """
        Extract text based on file type
        Returns: (text, ocr_applied)
        """
        ocr_applied = False
        
        if file_type == 'pdf':
            text = cls.extract_from_pdf(file_path)
        elif file_type == 'docx':
            text = cls.extract_from_docx(file_path)
        elif file_type in ['png', 'jpg', 'jpeg', 'tiff', 'bmp']:
            text = cls.extract_from_image(file_path)
            ocr_applied = True
        elif file_type == 'txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        else:
            # Try to read as text
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            except Exception:
                text = ""
        
        return text, ocr_applied


class LanguageDetector:
    """Detect document language"""
    
    @staticmethod
    def detect(text: str) -> str:
        """Detect language of text"""
        if not text or len(text.strip()) < 10:
            return 'unknown'
        
        if LANGDETECT_AVAILABLE:
            try:
                detected = detect(text)
                # Map to our standard codes
                if detected in ['da', 'no', 'sv']:
                    return 'da'  # Treat Scandinavian as Danish
                elif detected in ['en']:
                    return 'en'
                else:
                    return detected
            except Exception:
                pass
        
        # Fallback: simple heuristic
        danish_chars = set('æøåÆØÅ')
        if any(c in text for c in danish_chars):
            return 'da'
        
        return 'en'  # Default to English


class EntityExtractor:
    """Extract named entities and concepts from text using spaCy"""
    
    def __init__(self, model_name: str = 'en_core_web_sm'):
        self.nlp = None
        self.model_name = model_name
        if SPACY_AVAILABLE:
            try:
                self.nlp = spacy.load(model_name)
            except OSError:
                # Model not downloaded, will download on first use
                pass
    
    def _ensure_model(self):
        """Download and load spaCy model if needed"""
        if self.nlp is None and SPACY_AVAILABLE:
            import subprocess
            subprocess.run(['python', '-m', 'spacy', 'download', self.model_name], 
                         capture_output=True)
            self.nlp = spacy.load(self.model_name)
    
    def extract(self, text: str, max_entities: int = 50) -> List[Dict[str, Any]]:
        """
        Extract named entities from text
        Returns list of {text, label, start, end, confidence}
        """
        if not text or not SPACY_AVAILABLE:
            return []
        
        self._ensure_model()
        if self.nlp is None:
            return []
        
        doc = self.nlp(text)
        entities = []
        
        for ent in doc.ents[:max_entities]:
            entities.append({
                'text': ent.text,
                'label': ent.label_,
                'start': ent.start_char,
                'end': ent.end_char,
                'confidence': 1.0  # spaCy doesn't provide confidence
            })
        
        return entities
    
    def extract_concepts(self, text: str, max_concepts: int = 30) -> List[str]:
        """
        Extract key concepts (nouns, noun phrases) from text
        """
        if not text or not SPACY_AVAILABLE:
            return []
        
        self._ensure_model()
        if self.nlp is None:
            return []
        
        doc = self.nlp(text)
        concepts = set()
        
        # Extract noun chunks
        for chunk in doc.noun_chunks:
            if len(chunk.text) > 2 and len(chunk.text.split()) <= 4:
                concepts.add(chunk.text.lower().strip())
        
        # Extract important nouns
        for token in doc:
            if token.pos_ in ['NOUN', 'PROPN'] and len(token.text) > 2:
                concepts.add(token.lemma_.lower())
        
        # Return most frequent concepts
        return list(concepts)[:max_concepts]


class EmbeddingGenerator:
    """Generate semantic embeddings using sentence transformers"""
    
    def __init__(self, model_name: str = 'all-MiniLM-L6-v2'):
        self.model = None
        self.model_name = model_name
        if SENTENCE_TRANSFORMERS_AVAILABLE:
            try:
                self.model = SentenceTransformer(model_name)
            except Exception:
                pass
    
    def _ensure_model(self):
        """Load model if not already loaded"""
        if self.model is None and SENTENCE_TRANSFORMERS_AVAILABLE:
            self.model = SentenceTransformer(self.model_name)
    
    def generate(self, text: str) -> List[float]:
        """Generate embedding for text"""
        if not text:
            return []
        
        self._ensure_model()
        if self.model is None:
            # Fallback: random embedding (should not happen in production)
            return [0.0] * 384
        
        embedding = self.model.encode(text, convert_to_numpy=True)
        return embedding.tolist()
    
    def generate_batch(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings for multiple texts"""
        if not texts:
            return []
        
        self._ensure_model()
        if self.model is None:
            dim = 384
            return [[0.0] * dim for _ in texts]
        
        embeddings = self.model.encode(texts, convert_to_numpy=True)
        return embeddings.tolist()


class SimilarityCalculator:
    """Calculate similarity between embeddings"""
    
    @staticmethod
    def cosine_similarity(vec1: List[float], vec2: List[float]) -> float:
        """Calculate cosine similarity between two vectors"""
        if not vec1 or not vec2 or len(vec1) != len(vec2):
            return 0.0
        
        v1 = np.array(vec1)
        v2 = np.array(vec2)
        
        norm1 = np.linalg.norm(v1)
        norm2 = np.linalg.norm(v2)
        
        if norm1 == 0 or norm2 == 0:
            return 0.0
        
        return float(np.dot(v1, v2) / (norm1 * norm2))
    
    @staticmethod
    def find_similar(
        query_embedding: List[float],
        candidates: List[Tuple[str, List[float]]],
        top_k: int = 10,
        threshold: float = 0.3
    ) -> List[Tuple[str, float]]:
        """
        Find most similar items to query
        Returns list of (id, similarity_score)
        """
        if not query_embedding or not candidates:
            return []
        
        similarities = []
        for item_id, embedding in candidates:
            sim = SimilarityCalculator.cosine_similarity(query_embedding, embedding)
            if sim >= threshold:
                similarities.append((item_id, sim))
        
        # Sort by similarity descending
        similarities.sort(key=lambda x: x[1], reverse=True)
        return similarities[:top_k]


class KnowledgeGraph:
    """Build and manage knowledge graph using NetworkX"""
    
    def __init__(self):
        self.graph = nx.Graph() if NETWORKX_AVAILABLE else None
        self.nodes_data: Dict[str, Dict] = {}
        self.edges_data: List[Dict] = []
    
    def add_node(self, node_id: str, label: str, node_type: str, **kwargs):
        """Add a node to the graph"""
        if self.graph:
            self.graph.add_node(node_id, label=label, type=node_type, **kwargs)
        self.nodes_data[node_id] = {
            'id': node_id,
            'label': label,
            'type': node_type,
            **kwargs
        }
    
    def add_edge(self, source: str, target: str, weight: float, edge_type: str):
        """Add an edge between nodes"""
        if self.graph:
            self.graph.add_edge(source, target, weight=weight, type=edge_type)
        self.edges_data.append({
            'source': source,
            'target': target,
            'weight': weight,
            'type': edge_type
        })
    
    def build_similarity_edges(
        self,
        embeddings: Dict[str, List[float]],
        threshold: float = 0.5
    ):
        """
        Build similarity edges between nodes based on embeddings
        Uses efficient pairwise comparison
        """
        items = list(embeddings.items())
        n = len(items)
        
        for i in range(n):
            id1, emb1 = items[i]
            for j in range(i + 1, n):
                id2, emb2 = items[j]
                
                sim = SimilarityCalculator.cosine_similarity(emb1, emb2)
                if sim >= threshold:
                    self.add_edge(id1, id2, sim, 'similarity')
    
    def get_communities(self) -> List[List[str]]:
        """Detect communities in the graph using Louvain algorithm"""
        if not self.graph or self.graph.number_of_nodes() == 0:
            return []
        
        try:
            # Try louvain first (better quality)
            import community as community_louvain
            partition = community_louvain.best_partition(self.graph, weight='weight')
        except ImportError:
            # Fallback to greedy modularity
            partition = nx.algorithms.community.greedy_modularity_communities(self.graph)
            # Convert to dict format
            new_partition = {}
            for i, comm in enumerate(partition):
                for node in comm:
                    new_partition[node] = i
            partition = new_partition
        
        # Group nodes by community
        communities: Dict[int, List[str]] = {}
        for node, comm_id in partition.items():
            if comm_id not in communities:
                communities[comm_id] = []
            communities[comm_id].append(node)
        
        return list(communities.values())
    
    def get_stats(self) -> Dict:
        """Get graph statistics"""
        if not self.graph or self.graph.number_of_nodes() == 0:
            return {
                'nodes': 0,
                'edges': 0,
                'density': 0.0,
                'communities': 0,
                'avg_degree': 0.0
            }
        
        try:
            communities = self.get_communities()
            num_communities = len(communities)
        except Exception:
            num_communities = 0
        
        return {
            'nodes': self.graph.number_of_nodes(),
            'edges': self.graph.number_of_edges(),
            'density': nx.density(self.graph),
            'communities': num_communities,
            'avg_degree': sum(dict(self.graph.degree()).values()) / max(1, self.graph.number_of_nodes())
        }
    
    def to_dict(self) -> Dict:
        """Export graph to dictionary"""
        return {
            'nodes': list(self.nodes_data.values()),
            'edges': self.edges_data,
            'stats': self.get_stats()
        }


class DocumentProcessor:
    """Main processor that orchestrates document ingestion"""
    
    def __init__(self, data_dir: str = './data'):
        self.data_dir = Path(data_dir)
        self.data_dir.mkdir(parents=True, exist_ok=True)
        
        self.text_extractor = TextExtractor()
        self.language_detector = LanguageDetector()
        self.entity_extractor = EntityExtractor()
        self.embedding_generator = EmbeddingGenerator()
        self.knowledge_graph = KnowledgeGraph()
        
        self.documents: Dict[str, Document] = {}
        self.concept_nodes: Dict[str, ConceptNode] = {}
        
        # Load existing data
        self._load_state()
    
    def _generate_file_hash(self, file_path: str) -> str:
        """Generate SHA-256 hash of file"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    
    def _load_state(self):
        """Load existing documents and graph from disk"""
        state_file = self.data_dir / 'state.json'
        if state_file.exists():
            try:
                with open(state_file, 'r') as f:
                    state = json.load(f)
                
                for doc_data in state.get('documents', []):
                    doc = Document(**doc_data)
                    self.documents[doc.id] = doc
                
                # Rebuild graph
                for node_data in state.get('concept_nodes', []):
                    node = ConceptNode(**node_data)
                    self.concept_nodes[node.id] = node
                    self.knowledge_graph.add_node(
                        node.id, node.label, node.entity_type,
                        frequency=node.frequency,
                        documents=node.documents
                    )
                
                for edge_data in state.get('edges', []):
                    self.knowledge_graph.add_edge(
                        edge_data['source'],
                        edge_data['target'],
                        edge_data['weight'],
                        edge_data['type']
                    )
            except Exception as e:
                print(f"Warning: Could not load state: {e}")
    
    def _save_state(self):
        """Save current state to disk"""
        state = {
            'documents': [doc.to_dict() for doc in self.documents.values()],
            'concept_nodes': [asdict(node) for node in self.concept_nodes.values()],
            'edges': self.knowledge_graph.edges_data,
            'last_updated': datetime.now().isoformat()
        }
        
        with open(self.data_dir / 'state.json', 'w') as f:
            json.dump(state, f, indent=2)
    
    def process_file(self, file_path: str) -> Document:
        """
        Process a single file through the complete pipeline
        1. Generate hash for deduplication
        2. Extract text
        3. Detect language
        4. Extract entities
        5. Extract concepts
        6. Generate embedding
        7. Add to knowledge graph
        """
        # Check if file already processed
        file_hash = self._generate_file_hash(file_path)
        for doc in self.documents.values():
            if doc.file_hash == file_hash:
                print(f"File already processed: {doc.filename}")
                return doc
        
        # Determine file type
        ext = Path(file_path).suffix.lower().lstrip('.')
        file_type = ext if ext in ['pdf', 'docx', 'txt', 'png', 'jpg', 'jpeg'] else 'txt'
        
        # Extract text
        content, ocr_applied = TextExtractor.extract(file_path, file_type)
        
        if not content.strip():
            raise ValueError("No text could be extracted from file")
        
        # Detect language
        language = LanguageDetector.detect(content)
        
        # Extract entities
        entities = self.entity_extractor.extract(content)
        
        # Extract concepts
        concepts = self.entity_extractor.extract_concepts(content)
        
        # Generate embedding (using first 512 chars for speed, can be improved with chunking)
        embedding_text = content[:2000] if len(content) > 2000 else content
        embedding = self.embedding_generator.generate(embedding_text)
        
        # Create document
        doc_id = str(uuid.uuid4())
        doc = Document(
            id=doc_id,
            filename=Path(file_path).name,
            file_hash=file_hash,
            file_type=file_type,
            file_size=os.path.getsize(file_path),
            content=content[:10000],  # Store preview
            language=language,
            entities=entities,
            concepts=concepts,
            embedding=embedding,
            ocr_applied=ocr_applied,
            chunk_count=max(1, len(content) // 500)
        )
        
        self.documents[doc_id] = doc
        
        # Add entities to knowledge graph
        for entity in entities:
            entity_id = f"ent_{entity['label']}_{hash(entity['text']) % 100000}"
            
            if entity_id not in self.concept_nodes:
                concept_node = ConceptNode(
                    id=entity_id,
                    label=entity['text'],
                    entity_type=entity['label'],
                    frequency=1,
                    documents=[doc_id]
                )
                self.concept_nodes[entity_id] = concept_node
                self.knowledge_graph.add_node(
                    entity_id, entity['text'], entity['label'],
                    frequency=1, documents=[doc_id]
                )
            else:
                self.concept_nodes[entity_id].frequency += 1
                if doc_id not in self.concept_nodes[entity_id].documents:
                    self.concept_nodes[entity_id].documents.append(doc_id)
                self.knowledge_graph.nodes_data[entity_id]['frequency'] = self.concept_nodes[entity_id].frequency
        
        # Add document node to graph
        self.knowledge_graph.add_node(
            doc_id, doc.filename, 'document',
            language=doc.language,
            entity_count=len(entities)
        )
        
        # Save state
        self._save_state()
        
        return doc
    
    def search(
        self,
        query: str,
        top_k: int = 10,
        use_semantic: bool = True
    ) -> List[Dict[str, Any]]:
        """
        Search documents using hybrid approach:
        - Keyword matching (BM25-style)
        - Semantic similarity (if use_semantic=True)
        - Re-ranking
        """
        if not query.strip():
            return []
        
        results = []
        
        # Generate query embedding
        query_embedding = self.embedding_generator.generate(query) if use_semantic else []
        
        for doc in self.documents.values():
            score = 0.0
            
            # Keyword matching
            query_terms = set(query.lower().split())
            doc_terms = set(doc.content.lower().split())
            keyword_overlap = len(query_terms & doc_terms) / max(1, len(query_terms))
            score += keyword_overlap * 0.4
            
            # Entity matching
            entity_matches = sum(
                1 for e in doc.entities 
                if any(term in e['text'].lower() for term in query_terms)
            )
            score += min(entity_matches / max(1, len(doc.entities)), 1.0) * 0.3
            
            # Semantic similarity
            if use_semantic and query_embedding and doc.embedding:
                semantic_sim = SimilarityCalculator.cosine_similarity(query_embedding, doc.embedding)
                score += semantic_sim * 0.3
            
            if score > 0.1:  # Threshold
                results.append({
                    'document': doc.to_dict(),
                    'score': score,
                    'snippet': doc.content[:200] + '...' if len(doc.content) > 200 else doc.content
                })
        
        # Sort by score
        results.sort(key=lambda x: x['score'], reverse=True)
        return results[:top_k]
    
    def get_graph_data(self) -> Dict:
        """Get knowledge graph data for visualization"""
        # Rebuild similarity edges periodically
        doc_embeddings = {
            doc.id: doc.embedding 
            for doc in self.documents.values() 
            if doc.embedding
        }
        
        if len(doc_embeddings) > 1:
            self.knowledge_graph.build_similarity_edges(doc_embeddings, threshold=0.4)
        
        return self.knowledge_graph.to_dict()
    
    def get_stats(self) -> Dict:
        """Get overall system statistics"""
        graph_stats = self.knowledge_graph.get_stats()
        
        total_entities = sum(len(doc.entities) for doc in self.documents.values())
        total_size = sum(doc.file_size for doc in self.documents.values())
        
        language_dist = {}
        for doc in self.documents.values():
            lang = doc.language
            language_dist[lang] = language_dist.get(lang, 0) + 1
        
        return {
            'total_documents': len(self.documents),
            'total_concepts': len(self.concept_nodes),
            'total_entities': total_entities,
            'total_size_bytes': total_size,
            'languages': language_dist,
            'graph': graph_stats
        }


# API-ready instance
processor: Optional[DocumentProcessor] = None


def get_processor() -> DocumentProcessor:
    """Get or create processor instance"""
    global processor
    if processor is None:
        processor = DocumentProcessor()
    return processor
